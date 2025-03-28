import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple
import json
from tqdm import tqdm
from langchain.text_splitter import RecursiveCharacterTextSplitter
import re

@dataclass
class StyleChunk:
    content: str
    rule_type: str
    section: str
    examples: List[str]
    metadata: Dict
    embedding: Optional[List[float]] = None

    def to_dict(self):
        data = asdict(self)
        data['embedding'] = self.embedding.tolist() if self.embedding is not None else None
        return data

class StyleGuideProcessor:
    def __init__(self, embedding_model_name: str = "sentence-transformers/all-MiniLM-L6-v2"):
        self.model = SentenceTransformer(embedding_model_name)
        self.dimension = self.model.get_sentence_embedding_dimension()
        self.index = faiss.IndexFlatL2(self.dimension)
        self.chunks: List[StyleChunk] = []
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            separators=["\n\n", "\n", ".", " ", ""]
        )

    def extract_sections(self, text: str) -> List[Tuple[str, str]]:
        """Extract sections based on headings and formatting."""
        lines = text.split('\n')
        sections = []
        current_section = ""
        current_content = []
        
        heading_patterns = [
            r'^#+\s+(.+)$',  # Markdown style
            r'^[A-Z][^a-z]+[.:]\s*(.+)$',  # SECTION: style
            r'^\d+\.\s+(.+)$'  # Numbered sections
        ]
        
        for line in lines:
            is_heading = False
            for pattern in heading_patterns:
                match = re.match(pattern, line.strip())
                if match:
                    if current_section:
                        sections.append((current_section, '\n'.join(current_content).strip()))
                    current_section = match.group(1)
                    current_content = []
                    is_heading = True
                    break
            
            if not is_heading:
                current_content.append(line)
        
        if current_section:
            sections.append((current_section, '\n'.join(current_content).strip()))
        
        return sections

    def extract_examples(self, text: str) -> List[str]:
        """Extract examples from text content."""
        example_patterns = [
            r'(?:Example|e\.g\.|For example)[:\s]+([^.\n]+(?:[.\n][^.\n]+)*)',
            r'(?m)^•\s*([^.\n]+(?:[.\n][^.\n]+)*)',
            r'(?m)^\d+\.\s+([^.\n]+(?:[.\n][^.\n]+)*)'
        ]
        
        examples = []
        for pattern in example_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE)
            examples.extend(match.group(1).strip() for match in matches)
        
        return examples

    def identify_rule_type(self, text: str, section: str) -> str:
        """Identify rule type based on content and section."""
        rule_types = {
            'formatting': ['format', 'style', 'font', 'spacing', 'margin', 'indent', 'layout'],
            'grammar': ['grammar', 'tense', 'verb', 'noun', 'sentence', 'phrase'],
            'punctuation': ['punctuation', 'comma', 'period', 'colon', 'semicolon'],
            'terminology': ['term', 'word', 'vocabulary', 'glossary', 'definition'],
            'structure': ['structure', 'organization', 'section', 'heading', 'outline']
        }
        
        text_lower = f"{text.lower()} {section.lower()}"
        scores = {rule_type: 0 for rule_type in rule_types}
        
        for rule_type, keywords in rule_types.items():
            for keyword in keywords:
                if keyword in text_lower:
                    scores[rule_type] += 1
        
        max_score = max(scores.values())
        if max_score > 0:
            return max(scores.items(), key=lambda x: x[1])[0]
        return 'general'

    def preprocess_pdf(self, pdf_path: str) -> List[StyleChunk]:
        """Process PDF into semantic chunks efficiently."""
        chunks = []
        total_pages = 0
        current_page = 0
        
        # First pass to count pages
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            
        # Second pass to process pages
        with pdfplumber.open(pdf_path) as pdf:
            # Extract text with proper formatting
            full_text = ""
            for page in pdf.pages:
                current_page += 1
                if hasattr(self, 'progress_callback'):
                    self.progress_callback('style_guide', {
                        'phase': 'reading',
                        'current': current_page,
                        'total': total_pages,
                        'message': f'Reading page {current_page} of {total_pages}'
                    })
                    
                text = page.extract_text()
                if text:
                    full_text += text + "\n\n"
        
        # Extract sections
        sections = self.extract_sections(full_text)
        total_sections = len(sections)
        
        # Process each section
        for idx, (section_title, section_content) in enumerate(sections, 1):
            if hasattr(self, 'progress_callback'):
                self.progress_callback('style_guide', {
                    'phase': 'processing',
                    'current': idx,
                    'total': total_sections,
                    'message': f'Processing section {idx} of {total_sections}: {section_title}'
                })
                
            # Split section into smaller chunks
            text_chunks = self.text_splitter.split_text(section_content)
            
            for chunk in text_chunks:
                if len(chunk.strip()) < 50:  # Skip very small chunks
                    continue
                
                examples = self.extract_examples(chunk)
                rule_type = self.identify_rule_type(chunk, section_title)
                
                # Create chunk with embedding
                embedding = self.model.encode(chunk)
                
                style_chunk = StyleChunk(
                    content=chunk,
                    rule_type=rule_type,
                    section=section_title,
                    examples=examples,
                    metadata={
                        "source": "style_guide",
                        "confidence": 1.0
                    },
                    embedding=embedding
                )
                
                chunks.append(style_chunk)
                # Add to FAISS index
                self.index.add(np.array([embedding], dtype=np.float32))
        
        self.chunks = chunks
        return chunks

    def process_style_guide(self, pdf_path: str) -> List[StyleChunk]:
        """Process the style guide PDF and return the processed chunks."""
        return self.preprocess_pdf(pdf_path)

    def apply_style_rules(self, text: str, rules: List[Dict]) -> str:
        """Apply style rules to text and generate corrected version with change markers."""
        corrected_text = text
        has_changes = False
        
        # Sort rules by their position in the text (if they match)
        # This ensures we apply changes from end to start to maintain correct positions
        rules_with_pos = []
        for rule in rules:
            # Find the position where this rule would apply
            match_pos = corrected_text.lower().find(rule['rule'].lower())
            if match_pos >= 0:
                rules_with_pos.append((match_pos, rule))
        
        # Sort by position in reverse order
        rules_with_pos.sort(reverse=True)
        
        # Apply each rule
        for _, rule in rules_with_pos:
            # Find the text to replace
            match_pos = corrected_text.lower().find(rule['rule'].lower())
            if match_pos >= 0 and rule['examples']:
                # Get example to use as replacement
                replacement = rule['examples'][0]  # Use first example
                # Only mark as a change if the text actually differs
                original_text = corrected_text[match_pos:match_pos + len(rule['rule'])]
                if original_text.strip().lower() != replacement.strip().lower():
                    # Add change markers with confidence
                    marked_replacement = f"<change confidence={rule['confidence']}>{replacement}</change>"
                    # Replace the text
                    corrected_text = (
                        corrected_text[:match_pos] +
                        marked_replacement +
                        corrected_text[match_pos + len(rule['rule']):]
                    )
                    has_changes = True
        
        return corrected_text if has_changes else text

    def process_csr_document(self, doc_path: str) -> Tuple[Document, List[Dict]]:
        """Process CSR document and return both processed document and applied rules."""
        # Check if we have any style guide chunks
        if not self.chunks:
            raise ValueError("No style guide has been processed yet. Please upload a style guide first.")

        doc = Document(doc_path)
        processed_doc = Document()
        applied_rules = []
        
        total_paragraphs = len(doc.paragraphs)
        processed_paragraphs = 0
        
        for paragraph in doc.paragraphs:
            processed_paragraphs += 1
            if hasattr(self, 'progress_callback'):
                self.progress_callback('csr', {
                    'phase': 'processing',
                    'current': processed_paragraphs,
                    'total': total_paragraphs,
                    'message': f'Processing paragraph {processed_paragraphs} of {total_paragraphs}'
                })
                
            if not paragraph.text.strip():
                continue
            
            # Find relevant style rules
            embedding = self.model.encode(paragraph.text)
            k = min(3, len(self.chunks))  # Don't request more neighbors than we have chunks
            if k == 0:
                continue
                
            D, I = self.index.search(np.array([embedding], dtype=np.float32), k=k)
            
            # Get matched rules
            matched_rules = []
            for idx, distance in zip(I[0], D[0]):
                if idx >= 0 and idx < len(self.chunks) and distance < 1.5:  # Check index bounds and threshold
                    chunk = self.chunks[idx]
                    # Only include rules that have examples
                    if chunk.examples:
                        matched_rules.append({
                            "rule": chunk.content,
                            "type": chunk.rule_type,
                            "section": chunk.section,
                            "examples": chunk.examples,
                            "confidence": 1 - (distance / 2)  # Normalize distance to confidence
                        })
            
            # Apply rules and add to processed document
            new_para = processed_doc.add_paragraph(paragraph.text)
            
            # Generate corrected text if there are any rules
            if matched_rules:
                corrected_text = self.apply_style_rules(paragraph.text, matched_rules)
                # Only include in applied_rules if there were actual changes
                if corrected_text != paragraph.text:
                    applied_rules.append({
                        "original_text": paragraph.text,
                        "corrected_text": corrected_text,
                        "applied_rules": [rule for rule in matched_rules if any(
                            example.strip().lower() != paragraph.text.strip().lower()
                            for example in rule['examples']
                        )]
                    })
        
        return processed_doc, applied_rules

    def save_index(self, path: str):
        """Save FAISS index and chunks to disk."""
        faiss.write_index(self.index, f"{path}/style_index")
        with open(f"{path}/chunks.json", 'w') as f:
            json.dump([chunk.to_dict() for chunk in self.chunks], f)

    def load_index(self, path: str):
        """Load FAISS index and chunks from disk."""
        self.index = faiss.read_index(f"{path}/style_index")
        with open(f"{path}/chunks.json", 'r') as f:
            chunks_data = json.load(f)
            self.chunks = []
            for chunk_data in chunks_data:
                embedding = chunk_data.pop('embedding')
                chunk = StyleChunk(**chunk_data)
                chunk.embedding = np.array(embedding) if embedding else None
                self.chunks.append(chunk)

    def set_progress_callback(self, callback):
        """Set a callback function to receive progress updates."""
        self.progress_callback = callback
