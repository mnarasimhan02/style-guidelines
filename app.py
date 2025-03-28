from flask import Flask, request, jsonify, send_file, Response
import tempfile
import os
import json
from document_processor import StyleGuideProcessor
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_cors import CORS
from queue import Queue
from threading import Event
import io

app = Flask(__name__)
CORS(app)
processor = StyleGuideProcessor()

# Store progress updates for each client
progress_queues = {}

def progress_callback(client_id, progress_data):
    """Callback function to receive progress updates."""
    if client_id in progress_queues:
        progress_queues[client_id].put(progress_data)

@app.route('/api/progress/<client_id>')
def progress(client_id):
    """SSE endpoint for progress updates."""
    def generate():
        q = Queue()
        progress_queues[client_id] = q
        
        try:
            while True:
                progress_data = q.get()
                if progress_data == 'DONE':
                    break
                yield f"data: {json.dumps(progress_data)}\n\n"
        finally:
            progress_queues.pop(client_id, None)
    
    return Response(generate(), mimetype='text/event-stream')

@app.route('/')
def index():
    return jsonify({"message": "Style Guide Processor API is running"})

@app.route('/api/upload-style-guide', methods=['POST'])
def upload_style_guide():
    if 'style_guide' not in request.files:
        return jsonify({'success': False, 'error': 'No style guide file provided'}), 400
    
    file = request.files['style_guide']
    client_id = request.headers.get('X-Client-Id')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400

    if not file.filename.endswith('.pdf'):
        return jsonify({'success': False, 'error': 'Please upload a PDF file'}), 400

    try:
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, 'style_guide.pdf')
        file.save(temp_path)
        
        # Set up progress callback
        if client_id:
            processor.set_progress_callback(lambda doc_type, data: progress_callback(client_id, {
                'document': doc_type,
                **data
            }))
        
        chunks = processor.process_style_guide(temp_path)
        os.remove(temp_path)
        os.rmdir(temp_dir)
        
        # Signal completion
        if client_id in progress_queues:
            progress_queues[client_id].put('DONE')
        
        return jsonify({
            'success': True,
            'chunks': len(chunks) if chunks else 0
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/process-csr', methods=['POST'])
def process_csr():
    if 'csr_doc' not in request.files:
        return jsonify({'success': False, 'error': 'No CSR document provided'}), 400
    
    file = request.files['csr_doc']
    client_id = request.headers.get('X-Client-Id')
    
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400

    if not file.filename.endswith('.docx'):
        return jsonify({'success': False, 'error': 'Please upload a DOCX file'}), 400

    try:
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, 'csr.docx')
        file.save(temp_path)
        
        # Set up progress callback
        if client_id:
            processor.set_progress_callback(lambda doc_type, data: progress_callback(client_id, {
                'document': doc_type,
                **data
            }))
        
        processed_doc, applied_rules = processor.process_csr_document(temp_path)
        
        # Save the processed document to a temporary file
        output_path = os.path.join(temp_dir, 'processed_csr.docx')
        processed_doc.save(output_path)
        
        # Read the processed document content
        with open(output_path, 'rb') as f:
            doc_content = f.read()
        
        # Clean up temporary files
        os.remove(temp_path)
        os.remove(output_path)
        os.rmdir(temp_dir)
        
        # Signal completion
        if client_id in progress_queues:
            progress_queues[client_id].put('DONE')
        
        # Create a serializable response
        response_data = {
            'success': True,
            'results': {
                'applied_rules': applied_rules,
                'document_stats': {
                    'total_paragraphs': len(processed_doc.paragraphs),
                    'total_rules_applied': len(applied_rules)
                }
            }
        }
        
        return jsonify(response_data)
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

def create_corrected_doc(results):
    """Create a Word document with only corrected text."""
    doc = Document()
    
    # Add title
    doc.add_heading('Corrected Document', 0)
    
    # Add each corrected paragraph
    for item in results['applied_rules']:
        text = item['corrected_text'] if item['corrected_text'] else item['original_text']
        # Remove the XML-style tags for change markers
        clean_text = text.replace(
            '<change confidence=', ''
        ).replace('</change>', '').replace(
            '>', '', 1  # Remove first occurrence only (closing bracket of confidence value)
        )
        doc.add_paragraph(clean_text)
    
    return doc

def create_analysis_doc(results):
    """Create a Word document with corrections and applied rules."""
    doc = Document()
    
    # Add title
    doc.add_heading('Document Analysis', 0)
    
    # Add each paragraph with its analysis
    for idx, item in enumerate(results['applied_rules'], 1):
        # Section heading
        doc.add_heading(f'Paragraph {idx}', level=1)
        
        # Original text
        doc.add_heading('Original Text:', level=2)
        doc.add_paragraph(item['original_text'])
        
        # Corrected text
        doc.add_heading('Corrected Text:', level=2)
        if item['corrected_text']:
            clean_text = item['corrected_text'].replace(
                '<change confidence=', ''
            ).replace('</change>', '').replace(
                '>', '', 1
            )
            doc.add_paragraph(clean_text)
        else:
            doc.add_paragraph('No corrections needed')
        
        # Applied rules
        if item['applied_rules']:
            doc.add_heading('Applied Rules:', level=2)
            for rule in item['applied_rules']:
                p = doc.add_paragraph()
                p.add_run(f"Rule Type: ").bold = True
                p.add_run(f"{rule['type']}\n")
                p.add_run(f"Section: ").bold = True
                p.add_run(f"{rule['section']}\n")
                p.add_run(f"Rule: ").bold = True
                p.add_run(f"{rule['rule']}\n")
                if rule['examples']:
                    p.add_run(f"Examples:\n").bold = True
                    for example in rule['examples']:
                        doc.add_paragraph(example, style='List Bullet')
                doc.add_paragraph()  # Add space between rules
        
        # Add a page break between paragraphs
        doc.add_page_break()
    
    return doc

@app.route('/api/download-corrected', methods=['POST'])
def download_corrected():
    """Generate and download corrected text as a Word document."""
    try:
        results = request.json
        doc = create_corrected_doc(results)
        
        # Save to bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='corrected_document.docx'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/download-analysis', methods=['POST'])
def download_analysis():
    """Generate and download analysis as a Word document."""
    try:
        results = request.json
        doc = create_analysis_doc(results)
        
        # Save to bytes buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='document_analysis.docx'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8088, debug=True)
