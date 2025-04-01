from typing import List, Dict, Any
import PyPDF2
from docx import Document
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import re
import io
from app.models.rule_schema import StyleRule, RuleCategory
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.model = SentenceTransformer('paraphrase-MiniLM-L6-v2')
        self.index = None
        self.style_rules = []
        self.chunk_size = 500
        self.corrections_map = {
            # Clinical Terms (ordered by specificity)
            r'\btreatment[- ]emergent\s+adverse\s+event(?!\s*\([TEAE]+\))': r'treatment-emergent adverse event (TEAE)',
            r'\bserious\s+adverse\s+event(?!\s*\([SAE]+\))': r'serious adverse event (SAE)',
            r'\badverse\s+event(?!\s*\([AE]+\))': r'adverse event (AE)',
            r'\badverse\s+drug\s+reaction(?!\s*\([ADR]+\))': r'adverse drug reaction (ADR)',
            r'\binvestigational\s+product(?!\s*\([IP]+\))': r'investigational product (IP)',
            r'\bconcomitant\s+medication(?!\s*\([CM]+\))': r'concomitant medication (CM)',
            r'\bquality\s+of\s+life(?!\s*\([QoL]+\))': r'quality of life (QoL)',
            
            # Statistical Terms
            r'\bp-value\b': r'P value',
            r'\b(\d+%?\s*)ci\b': r'\1CI',
            r'\b(mean|median)\s*\((sd|se)\)': lambda m: f"{m.group(1)} ({m.group(2).upper()})",
            r'\bitt\b': r'ITT',
            r'\bpp\b': r'PP',
            r'\bodds\s+ratio(?!\s*\([OR]+\))': r'odds ratio (OR)',
            r'\bhazard\s+ratio(?!\s*\([HR]+\))': r'hazard ratio (HR)',
            r'\b(standard\s+error|mean|median)\s*\((sd|se)\)': lambda m: f"{m.group(1)} ({m.group(2).upper()})",
            
            # Units and Numbers
            r'(\d+)(\s*(?:mg|mL|L|kg|cm))\b': r'\1 \2',  # Add space between number and unit
            r'(\d+)\s*ml\b': r'\1 mL',
            r'(\d+)\s*l\b': r'\1 L',
            r'(\d+)\s*mg\b': r'\1 mg',
            r'(\d+)\s*kg\b': r'\1 kg',
            r'approximately\s+(\d+)': r'~\1',
            r'greater than or equal to\s*(\d+)': r'≥\1',
            r'less than or equal to\s*(\d+)': r'≤\1',
            
            # Document Structure
            r'\bsynopsis\b': r'Synopsis',
            r'\bappendix\s+([A-Za-z])\b': lambda m: f"Appendix {m.group(1).upper()}",
            r'\btable\s+(\d+)\b': lambda m: f"Table {m.group(1)}",
            r'\bfigure\s+(\d+)\b': lambda m: f"Figure {m.group(1)}",
            r'\bmaterials\s+and\s+methods\b': r'Materials and Methods',
            r'\bresults\s+and\s+discussion\b': r'Results and Discussion',
            
            # Study Phase
            r'\bphase\s*(1|one|i)\b': r'Phase 1',
            r'\bphase\s*(2|two|ii)\b': r'Phase 2',
            r'\bphase\s*(3|three|iii)\b': r'Phase 3',
            r'\bphase\s*(4|four|iv)\b': r'Phase 4',
            
            # Organizations and Regulatory
            r'\bfda\b': r'FDA',
            r'\bema\b': r'EMA',
            r'\birb\b': r'IRB',
            r'\biec\b': r'IEC',
            r'\bich\b': r'ICH',
            r'\bgcp\b': r'GCP',
            r'\bdaiichi\s+sankyo\b': r'Daiichi Sankyo',
            
            # Time Points
            r'\bbase-line\b': r'baseline',
            r'\bfollow\s+up\b': r'follow-up',
            r'\bend\s+of\s+treatment(?!\s*\([EOT]+\))': r'end of treatment (EOT)',
            r'\bend\s+of\s+study(?!\s*\([EOS]+\))': r'end of study (EOS)',
            r'\bscreening\s+period\b': r'Screening Period',
            r'\btreatment\s+period\b': r'Treatment Period',
            
            # Medical Terms
            r'\becg\b': r'ECG',
            r'\bmri\b': r'MRI',
            r'\bct\s+scan\b': r'CT scan',
            r'\bdna\b': r'DNA',
            r'\brna\b': r'RNA',
            r'\bpcr\b': r'PCR',
            
            # Demographics
            r'\bbmi\b': r'BMI',
            r'\bwhite\b': r'White',
            r'\bblack\b': r'Black',
            r'\basian\b': r'Asian',
            r'\bother\b': r'Other',
            r'\bmale\b': r'Male',
            r'\bfemale\b': r'Female',
            
            # Formatting
            r'i\.e\.\s*([a-z])': r'i.e., \1',
            r'e\.g\.\s*([a-z])': r'e.g., \1',
            r'vs\.\s*([a-z])': r'vs \1',
            r'etc\.\s*([a-z])': r'etc. \1',
        }
        
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file"""
        text = []
        try:
            doc = Document(file_path)
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
                    
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
                            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks by sentences"""
        # Split by sentence endings (.!?) but preserve the delimiter
        sentences = [s.strip() for s in re.split(r'([.!?])', text) if s.strip()]
        chunks = []
        current_chunk = []
        current_length = 0
        
        for i in range(0, len(sentences), 2):
            # Get the sentence and its delimiter if available
            sentence = sentences[i]
            delimiter = sentences[i + 1] if i + 1 < len(sentences) else ""
            full_sentence = sentence + delimiter
            
            sentence_length = len(full_sentence)
            if current_length + sentence_length > self.chunk_size:
                if current_chunk:
                    chunks.append("".join(current_chunk))
                current_chunk = [full_sentence]
                current_length = sentence_length
            else:
                current_chunk.append(full_sentence)
                current_length += sentence_length
        
        if current_chunk:
            chunks.append("".join(current_chunk))
        
        return chunks

    def apply_style_corrections(self, text: str) -> tuple[str, List[str]]:
        """Apply style corrections to the text."""
        corrections = []
        corrected_text = text
        
        # First pass: Apply non-clinical terms
        for pattern, replacement in self.corrections_map.items():
            if not any(kw in pattern for kw in ['adverse', 'event', 'reaction']):
                try:
                    if callable(replacement):
                        corrected_text = re.sub(pattern, replacement, corrected_text)
                    else:
                        compiled_pattern = re.compile(pattern, re.IGNORECASE)
                        corrected_text = compiled_pattern.sub(replacement, corrected_text)
                    
                    if corrected_text != text:
                        corrections.append(f"Applied rule: {pattern} -> {replacement}")
                        text = corrected_text
                except Exception as e:
                    logger.error(f"Error applying correction for {pattern}: {str(e)}")
                    continue
        
        # Second pass: Apply clinical terms in order of specificity
        def replace_clinical_term(match: re.Match) -> str:
            """Handle clinical term replacements with proper precedence."""
            text = match.group(0)
            text_lower = text.lower()
            
            if 'treatment emergent adverse event' in text_lower or 'treatment-emergent adverse event' in text_lower:
                return re.sub(r'treatment[- ]emergent\s+adverse\s+event', 'treatment-emergent adverse event (TEAE)', text, flags=re.IGNORECASE)
            elif 'serious adverse event' in text_lower:
                return re.sub(r'serious\s+adverse\s+event', 'serious adverse event (SAE)', text, flags=re.IGNORECASE)
            elif 'adverse event' in text_lower:
                return re.sub(r'adverse\s+event', 'adverse event (AE)', text, flags=re.IGNORECASE)
            elif 'adverse drug reaction' in text_lower:
                return re.sub(r'adverse\s+drug\s+reaction', 'adverse drug reaction (ADR)', text, flags=re.IGNORECASE)
            return text
        
        # Apply clinical term replacements
        clinical_pattern = r'\b(treatment[- ]emergent\s+adverse\s+event|serious\s+adverse\s+event|adverse\s+event|adverse\s+drug\s+reaction)\b'
        corrected_text = re.sub(clinical_pattern, replace_clinical_term, corrected_text, flags=re.IGNORECASE)
        
        # Post-process formatting
        corrected_text = re.sub(r'i\.e\.,?\s*(\d+)', r'i.e., \1', corrected_text)  # Fix i.e. with numbers
        corrected_text = re.sub(r'~\s+(\d+)', r'~\1', corrected_text)  # Fix spacing after ~
        corrected_text = re.sub(r'([≤≥])\s+(\d+)', r'\1\2', corrected_text)  # Fix spacing after ≤≥
        
        return corrected_text, corrections

    def process_style_guide(self, content: bytes, rules: Dict[str, List[StyleRule]]):
        """Process the style guide content and store the rules"""
        # Flatten the rules list
        flat_rules = []
        for category, rule_list in rules.items():
            flat_rules.extend(rule_list)
        
        self.style_rules = flat_rules
        
        # Create embeddings for semantic search
        rule_texts = []
        for rule in flat_rules:
            rule_texts.append(f"{rule.pattern} {rule.replacement}")
                
        if rule_texts:
            # Create embeddings
            embeddings = self.model.encode(rule_texts)
            
            # Initialize FAISS index
            dimension = embeddings.shape[1]
            self.index = faiss.IndexFlatL2(dimension)
            
            # Add embeddings to index
            self.index.add(embeddings.astype('float32'))
            
            logger.info(f"Initialized FAISS index with {len(rule_texts)} rules")

    def process_csr(self, file_path: str) -> List[Dict[str, Any]]:
        """Process CSR document and find style matches"""
        try:
            # Extract text from DOCX
            text = self.extract_text_from_docx(file_path)
            logger.info(f"Extracted {len(text)} characters from DOCX")
            
            # Split into chunks
            chunks = self.chunk_text(text)
            logger.info(f"Split text into {len(chunks)} chunks")
            
            results = []
            
            # Process each chunk
            for chunk in chunks:
                # Apply style corrections
                corrected_text, changes = self.apply_style_corrections(chunk)
                
                # Create embedding for the chunk
                chunk_embedding = self.model.encode([chunk])[0]
                
                # Find similar style rules
                matches = []
                if self.index is not None:
                    D, I = self.index.search(
                        chunk_embedding.reshape(1, -1).astype('float32'),
                        min(3, len(self.style_rules))
                    )
                    
                    matches = [
                        {
                            "rule": self.style_rules[idx],
                            "distance": float(dist),
                            "changes": changes
                        }
                        for dist, idx in zip(D[0], I[0])
                        if dist < 100  # Filter out very distant matches
                    ]
                
                # Include all chunks, even if no corrections were made
                results.append({
                    "text": chunk,
                    "corrected_text": corrected_text,
                    "matches": matches if matches else [],
                    "changes": changes
                })
            
            logger.info(f"Processed {len(chunks)} chunks, found corrections for {len([r for r in results if r['changes']])} chunks")
            return results
            
        except Exception as e:
            logger.error(f"Error processing CSR: {str(e)}")
            raise

    def process_document(self, file_path: str) -> tuple[str, List[str]]:
        """Process a document and return corrected text and changes made."""
        if file_path.lower().endswith('.docx'):
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file type")
            
        return self.apply_style_corrections(text)

    def process_and_show_changes(self, file_path: str) -> None:
        """Process a document and print the changes."""
        corrected_text, changes = self.process_document(file_path)
        
        print("\nProcessed Document Text:")
        print("=" * 80)
        print(corrected_text)
        print("\nChanges Made:")
        print("=" * 80)
        for change in changes:
            print(f"- {change}")
